from docx import Document
from pywebio.input import input, TEXT, textarea
from pywebio.output import put_text, put_loading, put_file
from pywebio.session import set_env
from pywebio import start_sevser
from os.path import abspath

def main():
    set_env(title='建造DOCX文档')
    doc_name = input("请输入你的文件名：", type=TEXT, required=True,
                     help_text="Submit：提交，确定;  Reset：重置;  示例： 六（4）班 张③;  温馨提示：无需加题目、后缀")
    doc_title = input("请输入标题：", type=TEXT, required=True,
                      help_text="Submit：提交，确定;  Reset：重置;  示例：当微风拂过我的脸；  温馨提示：不需要加书名号")
    doc_main = textarea("请输入正文：", minlength=5, maxlength=500000, required=True,
                        help_text="Submit：提交，确定;  Reset：重置;  示例：这是一段不少于五个字的文章；  温馨提示：要分好自然段，注意标点符号")
    put_loading(shape='border', color='primary')
    document = Document()
    document.add_heading(doc_title, 0)
    document.add_paragraph(doc_main)
    document.save('text2docx/files/' + doc_name + '  《' + doc_title + '》.docx')
    content = open('text2docx/files/' + doc_name + '  《' + doc_title + '》.docx', 'rb').read()
    put_file(doc_name + '  《' + doc_title + '》.docx', content, '文章生成完毕  点击我以下载')
    # put_link(name='点击下载文件',url="http://docxdownload.gtrees.tk/" + doc_name + '  《' + doc_title + '》.docx',new_window=False)
    put_text('文件生成完毕，请按按钮下载↑')
if __name__ == "__main__":
    start_sevser(main)
